"""
Resume Builder engine using python-docx to generate executive ATS-compliant resumes following gold-standard formats.
"""
import re
from pathlib import Path
from typing import Dict, Any, List, Union
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styles import (
    ResumeStyles,
    apply_page_setup,
    add_section_bottom_line,
    add_horizontal_borders,
    configure_right_tab_stop,
    set_run_font,
)


class ResumeBuilder:
    """Builder class to construct a robust, executive ATS-optimized resume document."""

    def __init__(self, font_family: str = ResumeStyles.FONT_FAMILY):
        self.font_family = font_family
        self.doc = docx.Document()
        apply_page_setup(self.doc)

    def _add_paragraph(
        self,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=Pt(0),
        space_after=Pt(0),
        line_spacing=ResumeStyles.LINE_SPACING
    ):
        """Helper to create a paragraph with standardized spacing."""
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        return p

    def _add_formatted_runs(
        self,
        paragraph,
        text: str,
        default_bold=False,
        default_italic=False,
        default_underline=False,
        size=ResumeStyles.BODY_SIZE
    ):
        """
        Parses inline markdown tokens (**bold**, *italic*) and adds formatted runs to a paragraph.
        """
        if not text:
            return

        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
        tokens = pattern.split(text)

        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**') and len(token) >= 4:
                inner_text = token[2:-2]
                run = paragraph.add_run(inner_text)
                set_run_font(
                    run,
                    font_name=self.font_family,
                    size=size,
                    bold=True,
                    italic=default_italic,
                    underline=default_underline
                )
            elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
                inner_text = token[1:-1]
                run = paragraph.add_run(inner_text)
                set_run_font(
                    run,
                    font_name=self.font_family,
                    size=size,
                    bold=default_bold,
                    italic=True,
                    underline=default_underline
                )
            else:
                run = paragraph.add_run(token)
                set_run_font(
                    run,
                    font_name=self.font_family,
                    size=size,
                    bold=default_bold,
                    italic=default_italic,
                    underline=default_underline
                )

    def add_header(self, header_data: Dict[str, Any]):
        """Render the candidate name, contact information bar, and highlighted target title."""
        if not header_data:
            return

        # 1. Candidate Full Name
        name = header_data.get("name", "").strip().upper()
        if name:
            p_name = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(0),
                space_after=Pt(1),
                line_spacing=1.0
            )
            run_name = p_name.add_run(name)
            set_run_font(
                run_name,
                font_name=self.font_family,
                size=ResumeStyles.NAME_SIZE,
                bold=True
            )

        # 2. Contact Information Bar
        contact = header_data.get("contact", {})
        contact_parts = []
        if isinstance(contact, dict):
            for field in ["location", "phone", "email", "linkedin", "github", "website", "portfolio"]:
                if contact.get(field):
                    contact_parts.append(str(contact[field]).strip())
        elif isinstance(contact, list):
            contact_parts = [str(c).strip() for c in contact if str(c).strip()]
        elif isinstance(contact, str) and contact.strip():
            contact_parts = [contact.strip()]

        if contact_parts:
            contact_line = " | ".join(contact_parts)
            p_contact = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(0),
                space_after=Pt(3),
                line_spacing=1.0
            )
            run_contact = p_contact.add_run(contact_line)
            set_run_font(
                run_contact,
                font_name=self.font_family,
                size=ResumeStyles.CONTACT_SIZE
            )

        # 3. Target Job Title (Framed with sleek top & bottom horizontal borders and proper padding)
        job_title = header_data.get("job_title") or header_data.get("title")
        if job_title:
            job_title_str = str(job_title).strip()
            p_title = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=ResumeStyles.TITLE_BEFORE,
                space_after=ResumeStyles.TITLE_AFTER,
                line_spacing=1.0
            )
            add_horizontal_borders(p_title, sz="10", space="4")
            display_title = job_title_str.upper() if len(job_title_str) <= 30 else job_title_str
            run_title = p_title.add_run(display_title)
            set_run_font(
                run_title,
                font_name=self.font_family,
                size=ResumeStyles.TITLE_SIZE,
                bold=True
            )

    def add_section_heading(self, heading_title: str):
        """Render an uppercase section heading with a clean bottom divider line."""
        p = self._add_paragraph(
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=ResumeStyles.SECTION_BEFORE,
            space_after=ResumeStyles.SECTION_AFTER,
            line_spacing=1.0
        )
        add_section_bottom_line(p, sz="6")
        run = p.add_run(heading_title.upper())
        set_run_font(
            run,
            font_name=self.font_family,
            size=ResumeStyles.SECTION_HEADING_SIZE,
            bold=True
        )

    def add_summary(self, summary_text: str):
        """Render the professional summary section."""
        if not summary_text or not summary_text.strip():
            return
        self.add_section_heading("PROFESSIONAL SUMMARY")
        p = self._add_paragraph(
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=Pt(1),
            space_after=Pt(2),
            line_spacing=ResumeStyles.LINE_SPACING
        )
        self._add_formatted_runs(p, summary_text.strip(), size=ResumeStyles.BODY_SIZE)

    def add_bullet_point(self, text: str):
        """
        Render a genuine bullet point item with authentic round bullet glyph and hanging indent.
        """
        if not text or not str(text).strip():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = ResumeStyles.BULLET_LEFT_INDENT
        p.paragraph_format.first_line_indent = ResumeStyles.BULLET_FIRST_LINE_INDENT
        p.paragraph_format.space_before = ResumeStyles.BULLET_BEFORE
        p.paragraph_format.space_after = ResumeStyles.BULLET_AFTER
        p.paragraph_format.line_spacing = ResumeStyles.LINE_SPACING

        # Round bullet glyph + tab
        r_bullet = p.add_run("•\t")
        set_run_font(r_bullet, font_name=self.font_family, size=ResumeStyles.BODY_SIZE)

        # Add formatted text
        self._add_formatted_runs(p, str(text).strip(), size=ResumeStyles.BODY_SIZE)

    def add_skills_section(self, skills_list: List[Dict[str, str]]):
        """Render the Technical Skills section with true hanging indent so wrapped lines align flush."""
        if not skills_list:
            return
        self.add_section_heading("TECHNICAL SKILLS")

        for item in skills_list:
            category = item.get("category", "").strip()
            skills = item.get("skills", "").strip()

            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = ResumeStyles.BULLET_LEFT_INDENT
            p.paragraph_format.first_line_indent = ResumeStyles.BULLET_FIRST_LINE_INDENT
            p.paragraph_format.space_before = ResumeStyles.BULLET_BEFORE
            p.paragraph_format.space_after = ResumeStyles.BULLET_AFTER
            p.paragraph_format.line_spacing = ResumeStyles.LINE_SPACING

            # Bullet glyph + tab for alignment
            r_bullet = p.add_run("•\t")
            set_run_font(r_bullet, font_name=self.font_family, size=ResumeStyles.BODY_SIZE)

            if category:
                r_cat = p.add_run(f"{category}: ")
                set_run_font(r_cat, font_name=self.font_family, size=ResumeStyles.BODY_SIZE, bold=True)

            self._add_formatted_runs(p, skills, size=ResumeStyles.BODY_SIZE)

    def add_experience_section(self, experience_list: List[Dict[str, Any]]):
        """Render the Professional Experience section following two-column company/role headers."""
        if not experience_list:
            return
        self.add_section_heading("WORK EXPERIENCE")

        for idx, exp in enumerate(experience_list):
            company = exp.get("company", "").strip()
            location = exp.get("location", "").strip()
            role = exp.get("role", "").strip()
            dates = exp.get("dates", "").strip()
            overview = exp.get("overview", "").strip()
            bullets = exp.get("bullets", [])

            # Header Line 1: Company (Bold Left) ... Dates (Bold Right)
            p_hdr1 = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=ResumeStyles.ENTRY_BEFORE if idx > 0 else Pt(1.5),
                space_after=Pt(0),
                line_spacing=1.0
            )
            configure_right_tab_stop(p_hdr1, position=Inches(7.5))

            r_comp = p_hdr1.add_run(company)
            set_run_font(
                r_comp,
                font_name=self.font_family,
                size=ResumeStyles.ENTRY_TITLE_SIZE,
                bold=True
            )

            if dates:
                p_hdr1.add_run("\t")
                r_dates = p_hdr1.add_run(dates)
                set_run_font(
                    r_dates,
                    font_name=self.font_family,
                    size=ResumeStyles.ENTRY_TITLE_SIZE,
                    bold=True
                )

            # Header Line 2: Role (Italic Left) ... Location (Italic Right)
            if role or location:
                p_hdr2 = self._add_paragraph(
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=Pt(0.5),
                    space_after=Pt(1.5),
                    line_spacing=1.0
                )
                configure_right_tab_stop(p_hdr2, position=Inches(7.5))

                if role:
                    r_role = p_hdr2.add_run(role)
                    set_run_font(
                        r_role,
                        font_name=self.font_family,
                        size=ResumeStyles.BODY_SIZE,
                        italic=True
                    )

                if location:
                    p_hdr2.add_run("\t")
                    r_loc = p_hdr2.add_run(location)
                    set_run_font(
                        r_loc,
                        font_name=self.font_family,
                        size=ResumeStyles.BODY_SIZE,
                        italic=True
                    )

            # Context / Overview line (if present)
            if overview:
                p_over = self._add_paragraph(
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_before=Pt(0.5),
                    space_after=Pt(1.5),
                    line_spacing=ResumeStyles.LINE_SPACING
                )
                self._add_formatted_runs(
                    p_over,
                    overview,
                    default_italic=True,
                    size=ResumeStyles.BODY_SIZE
                )

            # Bullets
            for b in bullets:
                self.add_bullet_point(b)

    def add_projects_section(self, projects_list: List[Dict[str, Any]]):
        """Render Key / Technical Projects section."""
        if not projects_list:
            return
        self.add_section_heading("TECHNICAL PROJECTS")

        for idx, proj in enumerate(projects_list):
            title = proj.get("title", "").strip()
            tools = proj.get("tools", "").strip()
            dates = proj.get("dates", "").strip()
            bullets = proj.get("bullets", [])

            # Project Header: Title (Bold) | Tools (Italic) ... Dates (Right Tab)
            p_hdr = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=ResumeStyles.ENTRY_BEFORE if idx > 0 else Pt(1.5),
                space_after=Pt(1),
                line_spacing=1.0
            )
            configure_right_tab_stop(p_hdr, position=Inches(7.5))

            r_title = p_hdr.add_run(title)
            set_run_font(
                r_title,
                font_name=self.font_family,
                size=ResumeStyles.ENTRY_TITLE_SIZE,
                bold=True
            )

            if tools:
                r_pipe = p_hdr.add_run(" | ")
                set_run_font(r_pipe, font_name=self.font_family, size=ResumeStyles.BODY_SIZE)
                r_tools = p_hdr.add_run(tools)
                set_run_font(
                    r_tools,
                    font_name=self.font_family,
                    size=ResumeStyles.BODY_SIZE,
                    italic=True
                )

            if dates:
                p_hdr.add_run("\t")
                r_date = p_hdr.add_run(dates)
                set_run_font(
                    r_date,
                    font_name=self.font_family,
                    size=ResumeStyles.BODY_SIZE,
                    italic=True
                )

            # Bullets
            for b in bullets:
                self.add_bullet_point(b)

    def add_education_and_certifications(self, edu_list: List[Dict[str, Any]]):
        """Render Education, Fellowships, and Certifications section."""
        if not edu_list:
            return
        self.add_section_heading("EDUCATION & CERTIFICATIONS")

        for idx, item in enumerate(edu_list):
            institution = item.get("institution", "").strip()
            title = item.get("title", "").strip()
            dates = item.get("dates", "").strip()
            credential_id = item.get("credential_id", "").strip()
            bullets = item.get("bullets", [])

            p_hdr = self._add_paragraph(
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=ResumeStyles.ENTRY_BEFORE if idx > 0 else Pt(1.5),
                space_after=Pt(0),
                line_spacing=1.0
            )
            configure_right_tab_stop(p_hdr, position=Inches(7.5))

            r_inst = p_hdr.add_run(institution)
            set_run_font(
                r_inst,
                font_name=self.font_family,
                size=ResumeStyles.ENTRY_TITLE_SIZE,
                bold=True
            )

            if dates:
                p_hdr.add_run("\t")
                r_date = p_hdr.add_run(dates)
                set_run_font(
                    r_date,
                    font_name=self.font_family,
                    size=ResumeStyles.ENTRY_TITLE_SIZE,
                    bold=True
                )

            if title or credential_id:
                p_sub = self._add_paragraph(
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=Pt(0),
                    space_after=Pt(1),
                    line_spacing=1.0
                )
                configure_right_tab_stop(p_sub, position=Inches(7.5))

                if title:
                    r_title = p_sub.add_run(title)
                    set_run_font(
                        r_title,
                        font_name=self.font_family,
                        size=ResumeStyles.BODY_SIZE,
                        italic=True
                    )

                if credential_id:
                    p_sub.add_run("\t")
                    r_cid = p_sub.add_run(f"Credential ID: {credential_id}")
                    set_run_font(
                        r_cid,
                        font_name=self.font_family,
                        size=ResumeStyles.BODY_SIZE,
                        italic=True
                    )

            for b in bullets:
                self.add_bullet_point(b)

    def build(self, data: Dict[str, Any]):
        """Construct the entire resume in executive ATS sequence from normalized data."""
        # 1. Header (Name, Contact, Target Job Title)
        if "header" in data and data["header"]:
            self.add_header(data["header"])

        # 2. Professional Summary (Power Profile)
        if "summary" in data and data["summary"]:
            self.add_summary(data["summary"])

        # 3. Technical Skills (Prominently categorized for ATS indexing)
        if "technical_skills" in data and data["technical_skills"]:
            self.add_skills_section(data["technical_skills"])

        # 4. Work Experience (Core reverse-chronological achievements)
        if "professional_experience" in data and data["professional_experience"]:
            self.add_experience_section(data["professional_experience"])

        # 5. Technical Projects
        if "key_projects" in data and data["key_projects"]:
            self.add_projects_section(data["key_projects"])

        # 6. Education & Certifications
        if "education_and_certifications" in data and data["education_and_certifications"]:
            self.add_education_and_certifications(data["education_and_certifications"])

    def save(self, output_path: Union[str, Path]) -> Path:
        """Save the generated Word document to file, handling Windows file locking gracefully."""
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            self.doc.save(str(path))
            print(f"[+] Resume successfully created at: {path}")
            return path
        except PermissionError:
            # Fallback if document is currently open and locked in Microsoft Word
            stem = path.stem
            ext = path.suffix
            fallback_path = path.parent / f"{stem}_new{ext}"
            try:
                self.doc.save(str(fallback_path))
                print(f"[!] Warning: '{path.name}' is open in another program (like Word).")
                print(f"[+] Saved successfully to fallback location: {fallback_path}")
                return fallback_path
            except Exception as e:
                raise PermissionError(
                    f"Could not save to '{path}' because it is open in Microsoft Word. "
                    f"Please close Word and run `python build_resume.py` again."
                ) from e
